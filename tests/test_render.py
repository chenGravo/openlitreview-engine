from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from openlitreview.render import (
    CHINESE_BODY_FONT,
    CHINESE_HEADING_FONT,
    _postprocess_docx,
)


def test_postprocess_docx_applies_chinese_review_layout(tmp_path: Path) -> None:
    path = tmp_path / "review.docx"
    document = Document()
    document.add_heading("测试文献综述", level=0)
    document.add_paragraph("正文段落。")
    document.add_heading("参考文献", level=2)
    document.add_paragraph("[1] Smith A. Example[J]. Test, 2024, 1(1): 1-2.")
    document.save(path)

    _postprocess_docx(path)

    result = Document(path)
    section = result.sections[0]
    assert abs(section.left_margin - Cm(3.0)) < 2_000
    assert abs(section.right_margin - Cm(2.5)) < 2_000
    normal_fonts = result.styles["Normal"]._element.rPr.rFonts
    title_fonts = result.styles["Title"]._element.rPr.rFonts
    assert normal_fonts.get(qn("w:eastAsia")) == CHINESE_BODY_FONT
    assert normal_fonts.get(qn("w:hint")) == "eastAsia"
    assert title_fonts.get(qn("w:eastAsia")) == CHINESE_HEADING_FONT
    assert title_fonts.get(qn("w:hint")) == "eastAsia"
    reference = result.paragraphs[-1]
    assert abs(reference.paragraph_format.left_indent - Cm(0.74)) < 2_000
    assert abs(reference.paragraph_format.first_line_indent - Cm(-0.74)) < 2_000


def test_pinned_citeproc_renderer_matches_gbt_fixture(tmp_path: Path) -> None:
    node = shutil.which("node")
    project = Path(__file__).resolve().parents[1]
    renderer = project / "node" / "render-citations.cjs"
    if not node or not (project / "node" / "node_modules" / "citeproc").is_dir():
        pytest.skip("Pinned citeproc-js runtime is not installed")
    output = tmp_path / "cited.md"
    subprocess.run(
        [
            node,
            str(renderer),
            str(project / "tests" / "fixtures" / "render_review.md"),
            str(project / "tests" / "fixtures" / "render_references.json"),
            str(
                project
                / "src"
                / "openlitreview"
                / "assets"
                / "gb-t-7714-2025-numeric-bilingual.csl"
            ),
            str(project / "src" / "openlitreview" / "assets"),
            str(output),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    rendered = output.read_text(encoding="utf-8")
    assert "## 参考文献" in rendered
    assert "[1]" in rendered and "[2]" in rendered
    assert "Smith A. Effects of an example intervention" in rendered
    assert "DOI:10.1000/example.2020" in rendered
    fixture = project / "tests" / "fixtures" / "render_references.json"
    assert len(json.loads(fixture.read_text())) == 2
